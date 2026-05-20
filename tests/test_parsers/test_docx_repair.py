"""docx_repair — pre-mammoth sanitizer for malformed OOXML.

Real-world trigger: WPS Office sometimes leaves dangling image
relationships (Target="../NULL" or pointing at deleted media). Word
and WPS open the file fine — they skip the broken image — but mammoth,
markitdown, and python-docx all raise KeyError on the zip lookup.
The whole fallback chain then returns empty content and the doc gets
misclassified as ``empty_content``.

These tests guard:
  * the standalone module: detects, repairs, and fails open
  * the DocxParser integration: a corrupted-rel docx now parses
    instead of returning empty
"""
from __future__ import annotations

import io
import re
import zipfile

import pytest
from PIL import Image


def _png_bytes(color=(200, 50, 100)) -> bytes:
    buf = io.BytesIO()
    Image.new("RGB", (50, 50), color).save(buf, format="PNG")
    return buf.getvalue()


def _build_clean_docx_with_image() -> bytes:
    """A genuine docx with a heading, body, and one embedded image.
    We use this as the substrate for the corruption tests because
    it gives us a real <a:blip> pointing at a real rel — easy to
    rewrite into a dangling rel."""
    from docx import Document

    doc = Document()
    doc.add_heading("WPS Corruption Test Doc", level=1)
    doc.add_paragraph("Body line one — this must survive sanitization.")
    doc.add_paragraph("Body line two — also must survive.")
    doc.add_picture(io.BytesIO(_png_bytes()))
    doc.add_paragraph("Body line three after the image.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _corrupt_image_rel_to_null(raw: bytes) -> bytes:
    """Take a valid docx and turn its image rel into a WPS-style
    dangling reference: change the Target to ``../NULL`` and delete
    the media file from the zip. The <a:blip r:embed=...> in
    document.xml is left intact (which is exactly the wild-state
    that crashes mammoth)."""
    bio = io.BytesIO(raw)
    with zipfile.ZipFile(bio) as zin:
        members = {item.filename: zin.read(item.filename) for item in zin.infolist()}
        infos = {item.filename: item for item in zin.infolist()}

    # Rewrite document.xml.rels: rewrite the image Target to "../NULL".
    rels = members["word/_rels/document.xml.rels"].decode("utf-8")
    rels_new = re.sub(
        r'(<Relationship\b[^/>]*Type="[^"]*/image"[^/>]*Target=")[^"]+(")',
        r"\1../NULL\2",
        rels,
        count=1,
    )
    assert rels_new != rels, "fixture builder didn't find an image rel to corrupt"
    members["word/_rels/document.xml.rels"] = rels_new.encode("utf-8")

    # Drop the media file the rel was pointing at — that's what the
    # WPS bug actually does.
    for name in list(members):
        if name.startswith("word/media/"):
            del members[name]

    bio_out = io.BytesIO()
    with zipfile.ZipFile(bio_out, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in members.items():
            info = infos.get(name) or zipfile.ZipInfo(name)
            zout.writestr(info, data)
    return bio_out.getvalue()


def _corrupt_image_crc(raw: bytes) -> bytes:
    """Take a valid docx with an embedded image and produce the WPS
    wild state: image payload bytes mutated but CRC in the zip header
    left pointing at the original. Reading the member back via
    ``zipfile`` then raises ``BadZipFile("Bad CRC-32 ...")``.

    Two-step build: first re-pack the image member with
    ``compress_type=ZIP_STORED`` so the on-disk bytes ARE the payload
    bytes (a DEFLATE-compressed payload would trip ``zlib.error`` on
    byte-flips before the CRC check ever ran — different failure
    surface than production). Then flip bytes in place, leaving the
    stored CRC untouched.
    """
    import struct

    # Step 1: re-pack with the image stored uncompressed.
    bio = io.BytesIO(raw)
    with zipfile.ZipFile(bio) as zin:
        items = [(i, zin.read(i.filename)) for i in zin.infolist()]
    bio_out = io.BytesIO()
    with zipfile.ZipFile(bio_out, "w") as zout:
        for item, data in items:
            info = zipfile.ZipInfo(item.filename, item.date_time)
            info.compress_type = (
                zipfile.ZIP_STORED if item.filename.startswith("word/media/")
                else zipfile.ZIP_DEFLATED
            )
            zout.writestr(info, data)
    repacked = bio_out.getvalue()

    # Step 2: find the stored image member's data offset and flip bytes.
    # LFH layout: 30 fixed bytes + filename + extra. Data follows.
    with zipfile.ZipFile(io.BytesIO(repacked)) as zf:
        media = [i for i in zf.infolist() if i.filename.startswith("word/media/")]
        assert media, "fixture builder needs at least one word/media/ entry"
        info = media[0]

    buf = bytearray(repacked)
    lfh_start = info.header_offset
    name_len = struct.unpack_from("<H", buf, lfh_start + 26)[0]
    extra_len = struct.unpack_from("<H", buf, lfh_start + 28)[0]
    data_start = lfh_start + 30 + name_len + extra_len
    for offset in range(data_start, min(data_start + 16, len(buf))):
        buf[offset] ^= 0xFF
    return bytes(buf)


# ---------------------------------------------------------------------------
# Standalone sanitize_ooxml_package
# ---------------------------------------------------------------------------

def test_sanitize_clean_docx_is_noop():
    """Clean docx returns the input bytes object unchanged and reports
    no actions. The cheap path matters — every healthy upload hits this."""
    from parsers.docx_repair import sanitize_ooxml_package

    raw = _build_clean_docx_with_image()
    out, actions = sanitize_ooxml_package(raw)

    assert actions == []
    # Same object identity is the intent (no-op short-circuits before
    # the rewrite loop) — confirms we're not paying for a copy.
    assert out is raw


def test_sanitize_garbage_bytes_fails_open():
    """Non-zip input — sanitizer must not raise. Returns input as-is
    so the parser still tries its own malformed-bytes handling."""
    from parsers.docx_repair import sanitize_ooxml_package

    out, actions = sanitize_ooxml_package(b"definitely not a docx")
    assert out == b"definitely not a docx"
    assert actions == []


def test_sanitize_detects_and_repairs_dangling_image_rel():
    """The WPS bug pattern: image rel Target="../NULL", media missing
    from zip. Sanitizer should drop the rel, strip the orphan blip,
    and report one action."""
    from parsers.docx_repair import sanitize_ooxml_package

    raw = _corrupt_image_rel_to_null(_build_clean_docx_with_image())
    out, actions = sanitize_ooxml_package(raw)

    assert len(actions) == 1
    a = actions[0]
    assert a.kind == "dangling_image_rel"
    assert a.rels_file == "word/_rels/document.xml.rels"
    assert a.target == "../NULL"
    assert a.rel_id.startswith("rId")

    # Verify the patched bytes are valid zip and the rel is gone.
    with zipfile.ZipFile(io.BytesIO(out)) as z:
        rels = z.read("word/_rels/document.xml.rels").decode("utf-8")
        assert "../NULL" not in rels
        # The orphan blip should be gone from document.xml — exact
        # element absence is what unblocks mammoth.
        doc = z.read("word/document.xml").decode("utf-8")
        assert f'r:embed="{a.rel_id}"' not in doc


def test_sanitize_detects_and_repairs_bad_crc_image():
    """WPS sometimes writes media with a bad CRC in the local header.
    Word/WPS skip CRC verification; Python's zipfile validates on
    .read() so mammoth crashes mid-image-extraction. Sanitizer should
    detect the bad-CRC read, drop the rel, strip the orphan blip,
    AND drop the corrupt member from the re-packed zip (so the next
    read attempt doesn't hit the same wall)."""
    from parsers.docx_repair import sanitize_ooxml_package

    raw = _corrupt_image_crc(_build_clean_docx_with_image())
    out, actions = sanitize_ooxml_package(raw)

    assert len(actions) == 1
    a = actions[0]
    assert a.kind == "bad_crc_image"
    assert a.rels_file == "word/_rels/document.xml.rels"
    assert a.rel_id.startswith("rId")

    with zipfile.ZipFile(io.BytesIO(out)) as z:
        # The rel is gone.
        rels = z.read("word/_rels/document.xml.rels").decode("utf-8")
        assert f'Id="{a.rel_id}"' not in rels
        # The orphan blip is gone from document.xml.
        doc = z.read("word/document.xml").decode("utf-8")
        assert f'r:embed="{a.rel_id}"' not in doc
        # The bad-CRC media member itself is gone from the zip — the
        # rel that referenced it is gone, so leaving the bytes would
        # be dead weight (and re-reading them would just re-raise).
        assert not any(n.startswith("word/media/") for n in z.namelist())


@pytest.mark.asyncio
async def test_corrupted_crc_docx_parses_after_repair():
    """End-to-end via DocxParser: the WPS bad-CRC pattern produces an
    empty parse pre-fix (mammoth crashes inside its image walker).
    Post-fix the body text comes out and ``repair_kinds`` carries
    ``bad_crc_image`` for the metric."""
    from parsers.docx_parser import DocxParser

    raw = _corrupt_image_crc(_build_clean_docx_with_image())

    # Sanity: raw bytes crash mammoth's image-extraction path. The
    # default image handler base64-inlines, so we install a minimal
    # one that reads the bytes — the read is what trips the CRC check.
    import mammoth
    def _read_and_discard(image):
        with image.open() as stream:
            stream.read()
        return {"src": "x", "alt": ""}
    with pytest.raises(zipfile.BadZipFile):
        mammoth.convert_to_markdown(
            io.BytesIO(raw),
            convert_image=mammoth.images.img_element(_read_and_discard),
        )

    parser = DocxParser()
    result = await parser.parse(raw, filename="badcrc.docx")

    assert "Body line one" in result.content
    assert "Body line two" in result.content
    assert "Body line three" in result.content
    assert result.image_refs == []
    assert result.repair_kinds == ["bad_crc_image"]


def test_sanitize_external_target_is_left_alone():
    """``TargetMode="External"`` rels (hyperlinks to web URLs, sometimes
    image URLs too) legitimately point outside the zip — they must NOT
    be flagged as dangling."""
    from parsers.docx_repair import sanitize_ooxml_package

    # Build a docx, then inject an external image rel manually.
    raw = _build_clean_docx_with_image()
    bio = io.BytesIO(raw)
    with zipfile.ZipFile(bio) as zin:
        members = {n: zin.read(n) for n in zin.namelist()}
        infos = {item.filename: item for item in zin.infolist()}

    rels = members["word/_rels/document.xml.rels"].decode("utf-8")
    rels_new = rels.replace(
        "</Relationships>",
        '<Relationship Id="rIdExt" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" '
        'Target="https://example.com/missing.png" TargetMode="External"/>'
        "</Relationships>",
    )
    members["word/_rels/document.xml.rels"] = rels_new.encode("utf-8")

    bio_out = io.BytesIO()
    with zipfile.ZipFile(bio_out, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in members.items():
            zout.writestr(infos.get(name) or zipfile.ZipInfo(name), data)

    out, actions = sanitize_ooxml_package(bio_out.getvalue())
    assert actions == []


# ---------------------------------------------------------------------------
# DocxParser integration
# ---------------------------------------------------------------------------

@pytest.fixture(autouse=True)
def _disable_isolation(monkeypatch):
    from config import get_settings
    monkeypatch.setattr(get_settings(), "parser_isolation", False)


@pytest.mark.asyncio
async def test_corrupted_docx_parses_after_repair():
    """Pre-repair: mammoth crashes and our parser swallows it →
    ParseResult(content=""). Post-repair: the same bytes flow through
    the parser, mammoth succeeds on the sanitized zip, body text is
    extracted, and no synthetic image_ref appears (the dangling blip
    was structurally removed, not redirected to a placeholder)."""
    from parsers.docx_parser import DocxParser

    raw = _corrupt_image_rel_to_null(_build_clean_docx_with_image())

    # Sanity: mammoth direct call on the corrupted file should crash —
    # if this assertion ever stops holding, the bug is gone upstream
    # and this whole module is candidate for deletion.
    import mammoth
    with pytest.raises(KeyError):
        mammoth.convert_to_markdown(io.BytesIO(raw))

    parser = DocxParser()
    result = await parser.parse(raw, filename="corrupted.docx")

    assert "Body line one" in result.content
    assert "Body line two" in result.content
    assert "Body line three" in result.content
    # The orphan blip is gone — there was no real image left to capture.
    assert result.image_refs == []
    # Repair kinds round-trip back so the parent process can drive
    # PARSE_REPAIRS_TOTAL — the counter lives outside the parser
    # subprocess pool and won't see in-subprocess increments.
    assert result.repair_kinds == ["dangling_image_rel"]


@pytest.mark.asyncio
async def test_clean_docx_carries_no_repair_kinds():
    """Healthy uploads must not drag a junk repair_kinds entry through
    the pipeline — the parent's metric increment loop counts non-empty
    lists, so a stray entry would inflate ops dashboards."""
    from parsers.docx_parser import DocxParser
    from tests.test_parsers.test_docx_parser import _build_docx_bytes

    parser = DocxParser()
    result = await parser.parse(_build_docx_bytes(), filename="clean.docx")
    assert result.repair_kinds == []


@pytest.mark.asyncio
async def test_repair_kinds_survive_post_repair_mammoth_failure():
    """If sanitization runs but mammoth still crashes downstream
    (another quirk we don't know about yet), the parser returns
    empty content — but ``repair_kinds`` must still carry the
    successful sanitizer hits so the metric records them. Without
    this we'd silently undercount producer quirks in production."""
    from parsers.base import ParseResult
    from parsers import docx_parser as docx_parser_mod

    raw = _corrupt_image_rel_to_null(_build_clean_docx_with_image())

    # Stub mammoth to raise after sanitization runs — simulates a
    # second, unknown malformation that the sanitizer doesn't yet
    # cover. The parser must still emit repair_kinds for what it
    # *did* fix.
    class _BoomMammoth:
        @staticmethod
        def convert_to_markdown(*a, **kw):
            raise RuntimeError("simulated post-repair mammoth crash")

        class images:
            @staticmethod
            def img_element(handler):
                return handler

    import sys
    real_mammoth = sys.modules.get("mammoth")
    sys.modules["mammoth"] = _BoomMammoth
    try:
        result = await docx_parser_mod.DocxParser().parse(
            raw, filename="boom.docx",
        )
    finally:
        if real_mammoth is not None:
            sys.modules["mammoth"] = real_mammoth
        else:
            sys.modules.pop("mammoth", None)

    assert isinstance(result, ParseResult)
    assert result.content == ""
    assert result.repair_kinds == ["dangling_image_rel"]


@pytest.mark.asyncio
async def test_metric_increments_in_parent_process():
    """End-to-end metric path: the corrupted-docx parse must bump
    ``PARSE_REPAIRS_TOTAL`` in the *parent* (caller) process. The
    parser's _parse_sync runs inside the isolation subprocess pool
    and any .inc() there would be invisible to the worker's metrics
    endpoint — this test guards that the round-trip through
    ParseResult.repair_kinds is wired up correctly."""
    from infra.metrics import PARSE_REPAIRS_TOTAL
    from parsers.docx_parser import DocxParser

    before = PARSE_REPAIRS_TOTAL.labels(
        engine="docx-native", kind="dangling_image_rel",
    )._value.get()

    raw = _corrupt_image_rel_to_null(_build_clean_docx_with_image())
    await DocxParser().parse(raw, filename="metric.docx")

    after = PARSE_REPAIRS_TOTAL.labels(
        engine="docx-native", kind="dangling_image_rel",
    )._value.get()
    assert after == before + 1
