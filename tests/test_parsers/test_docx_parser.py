"""DocxParser (mammoth + python-docx) — native alternative to docling
for .docx text extraction.

Guards:
  * heading / list / table round-trip via mammoth → markdown
  * inline images captured as ImageRef (not base64-inlined in content)
  * registry prefers "docx-native" over "docling" for .docx
  * title falls back to filename stem when the doc has no H1
"""
import io

import pytest
from PIL import Image


def _png_bytes(color=(200, 50, 100)) -> bytes:
    buf = io.BytesIO()
    Image.new("RGB", (100, 100), color).save(buf, format="PNG")
    return buf.getvalue()


def _build_docx_bytes(with_image: bool = False) -> bytes:
    """Construct a real .docx in memory with heading + body + optional image."""
    from docx import Document

    doc = Document()
    doc.add_heading("Main Title", level=1)
    doc.add_paragraph("First paragraph body text.")
    doc.add_heading("Section One", level=2)
    doc.add_paragraph("Bullet-style content (not really bulleted).")

    if with_image:
        img = io.BytesIO(_png_bytes())
        doc.add_picture(img)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture(autouse=True)
def _disable_isolation(monkeypatch):
    """Run parser in-process for tests (same as the repo-wide default in
    conftest, but duplicated here to be explicit)."""
    from config import get_settings
    monkeypatch.setattr(get_settings(), "parser_isolation", False)


# ---------------------------------------------------------------------------
# Engine metadata
# ---------------------------------------------------------------------------

def test_engine_name_and_supported_types():
    from parsers.docx_parser import DocxParser
    assert DocxParser.engine_name == "docx-native"
    assert ".docx" in DocxParser.supported_types
    assert ".doc" in DocxParser.supported_types


def test_is_available_with_deps_present():
    from parsers.docx_parser import DocxParser
    # mammoth + python-docx are pinned; tests run in the uv env.
    assert DocxParser.is_available()


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

@pytest.mark.asyncio
async def test_parse_docx_returns_markdown_with_headings():
    from parsers.docx_parser import DocxParser

    docx_bytes = _build_docx_bytes()
    parser = DocxParser()
    result = await parser.parse(docx_bytes, filename="doc.docx")

    assert "Main Title" in result.content
    # mammoth escapes punctuation in paragraph text (``.`` → ``\.``,
    # ``(`` → ``\(``) per CommonMark rules. Check on the unescaped
    # core so we don't couple the test to mammoth's escaping policy.
    assert "First paragraph body text" in result.content
    assert "Section One" in result.content
    # mammoth emits heading markers — structure matters for our
    # outline extractor and for the splitter's semantic chunk anchor.
    assert result.content.startswith("# Main Title")


@pytest.mark.asyncio
async def test_parse_title_from_first_heading():
    from parsers.docx_parser import DocxParser
    docx_bytes = _build_docx_bytes()
    parser = DocxParser()
    result = await parser.parse(docx_bytes, filename="deck.docx")
    # First H1 is "Main Title"; title should come from there, not the
    # filename stem.
    assert result.title == "Main Title"


@pytest.mark.asyncio
async def test_parse_title_falls_back_to_filename_when_no_heading():
    from parsers.docx_parser import DocxParser
    from docx import Document
    doc = Document()
    doc.add_paragraph("Plain body, no heading at all.")
    buf = io.BytesIO()
    doc.save(buf)

    parser = DocxParser()
    result = await parser.parse(buf.getvalue(), filename="unnamed.docx")
    assert result.title == "unnamed"


# ---------------------------------------------------------------------------
# Image extraction
# ---------------------------------------------------------------------------

@pytest.mark.asyncio
async def test_parse_captures_inline_images_as_refs():
    """Content must NOT inline base64 data — that would balloon
    content_len and break the max_content_chars gate. Instead, images
    land in ``image_refs`` with lazy bytes factories.
    """
    from parsers.docx_parser import DocxParser
    docx_bytes = _build_docx_bytes(with_image=True)
    parser = DocxParser()
    result = await parser.parse(docx_bytes, filename="withimg.docx")

    assert len(result.image_refs) == 1
    ref = result.image_refs[0]
    assert ref.name.startswith("image_0000") and ref.name.endswith(".png")
    # Round-trip through PIL to verify bytes are real PNG.
    with Image.open(io.BytesIO(ref.materialize())) as img:
        assert img.size == (100, 100)

    # The markdown contains a reference to the image name, not the
    # full base64 blob.
    assert "image_0000.png" in result.content
    # Sanity: no base64 data-URI leaked into content.
    assert "data:image" not in result.content


@pytest.mark.asyncio
async def test_parse_handles_malformed_bytes(monkeypatch):
    """Garbage that isn't a zip and isn't anything LibreOffice can
    sniff. Both the legacy-route and the mammoth fast path must end in
    empty content (not an unhandled exception) so the parse chain can
    fall through to the next engine.
    """
    from parsers.docx_parser import DocxParser
    # Force LibreOffice to fail so the test doesn't depend on whether
    # the binary is installed in the test environment.
    def _boom(*_a, **_kw):
        raise RuntimeError("libreoffice unavailable in test")
    monkeypatch.setattr(
        "parsers.docx_parser.convert_legacy_format", _boom,
    )
    parser = DocxParser()
    result = await parser.parse(b"not a docx", filename="broken.docx")
    assert result.content == ""
    assert result.image_refs == []


@pytest.mark.asyncio
async def test_docx_filename_with_non_zip_content_routes_to_libreoffice(monkeypatch):
    """Reproduces a real prod failure: users upload .doc binary renamed
    .docx (or .wps, etc.). mammoth raises ``BadZipFile`` and the whole
    fallback chain silently produces empty_content. The magic-byte
    sniff must reroute these through LibreOffice so they parse.
    """
    from parsers.docx_parser import DocxParser
    # Non-zip payload — pretend it's a legacy .doc (OLE2 signature).
    fake_doc_bytes = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"x" * 64

    # Capture which suffix the parser hands to convert_legacy_format,
    # and return a valid .docx so mammoth has something real to chew.
    seen: dict = {}
    def _fake_convert(src_path: str, suffix: str) -> str:
        seen["suffix"] = suffix
        seen["src_path"] = src_path
        converted = src_path.replace(suffix, ".docx")
        with open(converted, "wb") as f:
            f.write(_build_docx_bytes())
        return converted
    monkeypatch.setattr(
        "parsers.docx_parser.convert_legacy_format", _fake_convert,
    )

    parser = DocxParser()
    result = await parser.parse(fake_doc_bytes, filename="renamed.docx")

    # Routed through libreoffice path, with the .doc suffix hint we
    # picked when sniffing said "not a zip".
    assert seen.get("suffix") == ".doc"
    # And the converted .docx parsed into real content.
    assert "Main Title" in result.content


# ---------------------------------------------------------------------------
# Registry dispatch
# ---------------------------------------------------------------------------

def test_registry_picks_docx_native_over_docling():
    """Ingest path routes .docx to mammoth, not docling. If docling is
    later removed entirely, this selection is what keeps things
    working.
    """
    from parsers.registry import registry
    # Force fresh evaluation (is_available is a classmethod, no caching).
    engine = registry.auto_select(filename="report.docx")
    assert engine == "docx-native", (
        f"expected docx-native, got {engine!r}; "
        "_EXT_PREFERENCE ordering probably regressed"
    )
